"""Exports compliant with Order #440 (МОУ, edition 25.04.2025).

Generates four forms used by ЗСУ for electronic military property accounting:
- Summary statement (Додаток 1) — inventory snapshot per item type
- Movement journal — chronological event log for a date range
- Loss act — irretrievable loss write-off (one per UsageEvent)
- Repair act — return to repair (one per UsageEvent)

The reference Order #440 structure is documented in docs/normative/mod-440.md.
Cells / fields kept here are the mandatory minimum; specific воинская часть can
extend via admin-side templates without code changes.
"""
from __future__ import annotations

from collections.abc import Iterable
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Cm
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter


@dataclass(frozen=True)
class MovementRow:
    event_date: date
    serial_no: str
    item_type_code: str
    operator_code: str
    outcome: str
    reason_code: str | None
    basis: str  # підстава: «акт від ДД.ММ.РРРР»


@dataclass(frozen=True)
class InventoryRow:
    item_type_code: str
    item_type_name: str
    measure_unit: str  # одиниця обліку, напр. «шт.»
    on_hand: int
    received: int
    consumed: int
    lost: int
    in_repair: int


@dataclass(frozen=True)
class LossActData:
    act_no: str
    act_date: date
    unit_name: str  # найменування військової частини
    serial_no: str
    item_type_name: str
    measure_unit: str
    operator_code: str
    event_date: date
    reason_code: str
    reason_name: str
    circumstances: str
    responsible_person: str


@dataclass(frozen=True)
class RepairActData:
    act_no: str
    act_date: date
    unit_name: str
    serial_no: str
    item_type_name: str
    operator_code: str
    event_date: date
    reason_code: str
    reason_name: str
    defect_description: str
    sender: str   # хто здав
    receiver: str  # хто прийняв на ремонт


_BOLD = Font(bold=True)
_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)


def _header(ws, title: str, *, span: int = 8) -> None:
    ws.append([title])
    cell = ws.cell(row=ws.max_row, column=1)
    cell.font = Font(bold=True, size=12)
    ws.merge_cells(start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=span)
    cell.alignment = _CENTER


def _table_header(ws, headers: list[str]) -> None:
    ws.append(headers)
    for c in ws[ws.max_row]:
        c.font = _BOLD
        c.alignment = _CENTER


def inventory_summary_xlsx(
    rows: Iterable[InventoryRow], *, unit_name: str, as_of: date
) -> bytes:
    """Додаток 1 до п. 7 розділу II — узагальнююча відомість обліку."""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Додаток 1"

    _header(
        ws,
        f"УЗАГАЛЬНЮЮЧА ВІДОМІСТЬ обліку військового майна\n"
        f"{unit_name} · станом на {as_of.isoformat()}",
        span=9,
    )
    ws.append([])
    _table_header(
        ws,
        [
            "№ з/п",
            "Найменування",
            "Код типу",
            "Од. обліку",
            "Залишок",
            "Надійшло",
            "Спожито",
            "Втрачено",
            "У ремонті",
        ],
    )
    for i, r in enumerate(rows, start=1):
        ws.append(
            [
                i,
                r.item_type_name,
                r.item_type_code,
                r.measure_unit,
                r.on_hand,
                r.received,
                r.consumed,
                r.lost,
                r.in_repair,
            ]
        )

    ws.append([])
    ws.append(
        ["", "Матеріально відповідальна особа", "", "", "",
         "_____________ /__________________/"]
    )

    for col, width in enumerate([6, 36, 12, 12, 12, 12, 12, 12, 12], start=1):
        ws.column_dimensions[get_column_letter(col)].width = width

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def movement_journal_xlsx(
    rows: Iterable[MovementRow], *, unit_name: str, date_from: date, date_to: date
) -> bytes:
    """Журнал руху військового майна за період."""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Журнал руху"

    _header(
        ws,
        f"ЖУРНАЛ РУХУ ВІЙСЬКОВОГО МАЙНА\n"
        f"{unit_name} · {date_from.isoformat()} — {date_to.isoformat()}",
        span=7,
    )
    ws.append([])
    _table_header(
        ws,
        [
            "№ з/п",
            "Дата операції",
            "Серійний №",
            "Тип",
            "Експлуатант",
            "Результат",
            "Код причини",
            "Підстава",
        ],
    )
    for i, r in enumerate(rows, start=1):
        ws.append(
            [
                i,
                r.event_date.isoformat(),
                r.serial_no,
                r.item_type_code,
                r.operator_code,
                r.outcome,
                r.reason_code or "-",
                r.basis,
            ]
        )

    for col, width in enumerate([6, 14, 18, 8, 12, 12, 14, 36], start=1):
        ws.column_dimensions[get_column_letter(col)].width = width

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _bold_p(doc: Any, text: str, *, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True


def loss_act_docx(data: LossActData) -> bytes:
    """Акт списання виробу (безповоротна втрата)."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _bold_p(doc, f"ЗАТВЕРДЖУЮ\nКомандир {data.unit_name}", center=False)
    doc.add_paragraph("__________________________ / __________________ /")
    doc.add_paragraph(f'"{data.act_date.day:02d}" {data.act_date.strftime("%m.%Y")}')

    _bold_p(doc, f"АКТ № {data.act_no}", center=True)
    _bold_p(doc, "про безповоротну втрату виробу", center=True)
    doc.add_paragraph(f"від {data.act_date.isoformat()}")

    doc.add_paragraph(
        f"Комісія {data.unit_name} склала цей акт у тому, що "
        f"виріб «{data.item_type_name}», серійний № {data.serial_no} "
        f"({data.measure_unit}), закріплений за експлуатантом "
        f"{data.operator_code}, був втрачений безповоротно "
        f"{data.event_date.isoformat()}."
    )
    doc.add_paragraph(
        f"Код причини: {data.reason_code} — {data.reason_name}."
    )
    doc.add_paragraph(f"Обставини: {data.circumstances}")
    doc.add_paragraph(
        f"Матеріально відповідальна особа: {data.responsible_person}."
    )

    doc.add_paragraph("\nЧлени комісії:")
    for _ in range(3):
        doc.add_paragraph("________________ / ________________________ /")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def repair_act_docx(data: RepairActData) -> bytes:
    """Акт повернення виробу в ремонт."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _bold_p(doc, f"АКТ № {data.act_no}", center=True)
    _bold_p(doc, "приймання-передачі виробу в ремонт", center=True)
    doc.add_paragraph(f"від {data.act_date.isoformat()}")
    doc.add_paragraph(f"Військова частина: {data.unit_name}")

    doc.add_paragraph(
        f"Здавальник ({data.operator_code}) передав, "
        f"а приймальник прийняв на ремонт виріб "
        f"«{data.item_type_name}», серійний № {data.serial_no}, "
        f"повернутий з експлуатації {data.event_date.isoformat()}."
    )
    doc.add_paragraph(
        f"Код причини повернення: {data.reason_code} — {data.reason_name}."
    )
    doc.add_paragraph(f"Опис дефекту: {data.defect_description}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Здав: {data.sender}")
    doc.add_paragraph("________________ / ________________________ /")
    doc.add_paragraph(f"Прийняв: {data.receiver}")
    doc.add_paragraph("________________ / ________________________ /")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
